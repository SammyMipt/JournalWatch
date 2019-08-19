from django.shortcuts import reverse
from django.views.generic import ListView, CreateView
from django.http import HttpResponse

import docx
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import requests
from datetime import datetime, timedelta
from io import BytesIO

from .models import Article
from .forms import AddingForm
from .themes import THEME_CHOICES


class AddArticle(CreateView):
    model = Article
    template_name = 'article/add_article.html'
    form_class = AddingForm
    success_url = 'add_article'

    def form_valid(self, form):
        form.instance.author = self.request.user
        return super(AddArticle, self).form_valid(form)

    def get_success_url(self):
        return reverse(self.success_url)


class GetArticles(ListView):
    model = Article
    template_name = 'article/get_article.html'
    success_url = 'get_article'

    def get_success_url(self):
        return reverse(self.success_url)


def add_head(paragraph, article):
    add_hyperlink(paragraph, article.article_url, 'Link', '0000FF', True)
    paragraph.add_run(" {0}".format(article.description))


def add_hyperlink(paragraph, url, text, color, underline):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    new_run = docx.oxml.shared.OxmlElement('w:r')

    rpr_elem = docx.oxml.shared.OxmlElement('w:rPr')

    if not color is None:
        color_elem = docx.oxml.shared.OxmlElement('w:color')
        color_elem.set(docx.oxml.shared.qn('w:val'), color)
        rpr_elem.append(color_elem)

    if not underline:
        u_elem = docx.oxml.shared.OxmlElement('w:u')
        u_elem.set(docx.oxml.shared.qn('w:val'), 'none')
        rpr_elem.append(u_elem)

    new_run.append(rpr_elem)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    insert_text = OxmlElement('w:instrText')
    insert_text.set(qn('xml:space'), 'preserve')
    insert_text.text = 'TOC \\o "1-2" \\h \\z \\u'

    fld_char_2 = OxmlElement('w:fldChar')
    fld_char_2.set(qn('w:fldCharType'), 'separate')
    fld_char_3 = OxmlElement('w:t')
    fld_char_3.text = "Right-click to update field."
    fld_char_2.append(fld_char_3)

    fld_char4 = OxmlElement('w:fldChar')
    fld_char4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fld_char)
    r_element.append(insert_text)
    r_element.append(fld_char_2)
    r_element.append(fld_char4)


def get_docx(request):
    document = Document()
    start_date = datetime.strptime(str(request.GET.get("start")), "%Y-%m-%d")
    end_date = datetime.strptime(str(request.GET.get("end")), "%Y-%m-%d") + timedelta(days=1)

    all_articles = Article.objects.all().filter(created_at__range=(start_date, end_date))
    articles = dict()
    for theme_choice in THEME_CHOICES:
        articles[theme_choice[1]] = all_articles.all().filter(theme=theme_choice[0])

    document.add_paragraph('Table of Contents', style='TOC Heading')
    paragraph = document.add_paragraph()
    add_toc(paragraph)

    document.add_page_break()

    for theme in articles.keys():
        if len(articles[theme]) > 0:
            document.add_heading(theme, level=1)
            for article in articles[theme]:
                document.add_heading(article.title, level=2)

                p = document.add_paragraph()
                add_head(p, article)

                if article.image_url:
                    image_response = requests.get(article.image_url)
                    binary_img = BytesIO(image_response.content)
                    document.add_picture(binary_img, width=Inches(5.5))

                document.add_paragraph(article.abstract)

                if article.keywords:
                    document.add_paragraph(article.keywords)

                document.add_page_break()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=download.docx'
    document.save(response)

    return response
